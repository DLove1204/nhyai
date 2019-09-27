from django.db.models.functions import TruncMonth
from django.db.models import Count
import datetime
from .ienum import FILETYPE
import subprocess
from .pdfreader import PdfReader
from django.contrib.auth.models import User, Group
from rest_framework import viewsets, views
from api.serializers import UserSerializer, GroupSerializer
from rest_framework.response import Response
from rest_framework import status
from rest_framework.exceptions import ParseError
# uload package
from rest_framework.parsers import MultiPartParser, FormParser
from .serializers import FileUploadSerializer, WordRecognitionSerializer, FileImageTerrorismUploadSerializer, FileVisionPornUploadSerializer, AudioFileUploadSerializer
from .models import FileUpload, WordRecognition, FileImageTerrorismUpload, FileVisionPornUpload
# Handle Image
from PIL import Image
from io import BytesIO
import json
from .video import check_video
from .ocr.chineseocr import OCR
from violentsurveillance.image_terrorism import image_terrorism
from violentsurveillance.vision_porn import vision_porn
from django.conf import settings
from .serializers import VideoFileUploadSerializer, OcrGeneralSerializer, OcrIDCardSerializer, AudioFileInspectionSerializer, ImageFileUploadSerializer, WordRecognitionInspectionSerializer, OcrDrivinglicenseSerializer, OcrVehiclelicenseSerializer, OcrBusinesslicenseSerializer, OcrBankcardSerializer, OcrHandWrittenSerializer, OcrVehicleplateSerializer, HistoryRecordSerializer
from .models import VideoFileUpload, AudioFileUpload, OcrGeneral, OcrIDCard, AudioFileInspection, ImageFileUpload, WordRecognitionInspection, OcrDrivinglicense, OcrVehiclelicense, OcrBusinesslicense, OcrBankcard, OcrHandWritten, OcrVehicleplate, HistoryRecord
import os
import shutil
import uuid
import cv2
from .kaldi.audios import audio
from .sensitives.sensitives import sensitiveClass
import wave
import contextlib
import codecs
import chardet
from django.core.files import File
from urllib.request import urlopen
from tempfile import NamedTemporaryFile
from pydub import AudioSegment
import filetype
import docx
from .filetype import FileType
import platform
if(platform.system() == "Windows"):
    import win32com.client as wc
    import pythoncom


def get_two_float(f_str, n):
    f_str = str(f_str)      # f_str = '{}'.format(f_str) 也可以转换为字符串
    a, b, c = f_str.partition('.')
    c = (c+"0"*n)[:n]       # 如论传入的函数有几位小数，在字符串后面都添加n为小数0
    return ".".join([a, c])


def UpdateHistoryRecord(serializer, filetype, result, maxtype, violence, porn):
    file_id = serializer.id
    file_name = serializer.image.name.split('/')[1]
    file_url = settings.FILE_URL + serializer.image.url
    file_type = filetype
    inspection_result = result

    violence_percent = "0"
    violence_sensitivity_level = "0"
    if violence is not None:
        violence_percent = get_two_float(float(violence) * 100, 2)
        if (float(violence) < 0.5):
            violence_sensitivity_level = "0"
        if (float(violence) >= 0.5 and float(violence) <= 0.9):
            violence_sensitivity_level = "1"
        if (float(violence) > 0.9):
            violence_sensitivity_level = "2"

    porn_percent = "0"
    porn_sensitivity_level = "0"
    if porn is not None:
        porn_percent = get_two_float(float(porn) * 100, 2)
        if (float(porn) < 0.5):
            porn_sensitivity_level = "0"
        if (float(porn) >= 0.5 and float(porn) <= 0.9):
            porn_sensitivity_level = "1"
        if (float(porn) > 0.9):
            porn_sensitivity_level = "2"

    max_sensitivity_type = maxtype
    max_sensitivity_level = violence_sensitivity_level
    process_status = 2
    system_id = serializer.system_id
    channel_id = serializer.channel_id
    user_id = serializer.user_id

    HistoryRecord.objects.create(
        file_id=file_id, file_name=file_name,
        file_url=file_url, file_type=file_type,
        inspection_result=inspection_result, max_sensitivity_type=max_sensitivity_type,
        max_sensitivity_level=max_sensitivity_level, violence_percent=violence_percent,
        violence_sensitivity_level=violence_sensitivity_level, porn_percent=porn_percent,
        porn_sensitivity_level=porn_sensitivity_level, process_status=process_status,
        system_id=system_id, channel_id=channel_id, user_id=user_id
    )


def RunShellWithReturnCode(command):
    p = subprocess.Popen(command, stdout=subprocess.PIPE,
                         stderr=subprocess.PIPE, shell=True)
    p.wait()
    output = ""
    error = ""
    while True:
        line = p.stdout.read()
        if not line:
            break
        output += line.decode("utf-8")

    while True:
        err = p.stderr.read()
        if not err:
            break
        error += err.decode("utf-8")
    return output, error


class UserViewSet(viewsets.ModelViewSet):
    """
    API endpoint that allows users to be viewed or edited.
    """
    queryset = User.objects.all().order_by('-date_joined')
    serializer_class = UserSerializer


class GroupViewSet(viewsets.ModelViewSet):
    """
    API endpoint that allows groups to be viewed or edited.
    """
    queryset = Group.objects.all()
    serializer_class = GroupSerializer


class FileUploadViewSet(viewsets.ModelViewSet):

    queryset = FileUpload.objects.all()
    serializer_class = FileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        # file_obj = self.request.data.get('datafile')
        # print (file_obj)
        # ...
        # do some stuff with uploaded file
        # ...
        # try:
        #     img = Image.open(file_obj)
        #     # img.verify()
        #     pic_io = BytesIO()
        #     img.save(pic_io,img.format)

        # except:
        #     raise ParseError("Unsupported image type")

        file_path = iserializer.datafile.path
        check_result = settings.VIOLENCE.check_violence(file_path)
        # print (check_result)

        result = {
            "ret": 0,
            "msg": "ok",
            "data": {
                "tag_list": [
                    {
                        "tag_name": "protest",
                        "probability": check_result['protest']
                    },
                    {
                        "tag_name": "violence",
                        "probability": check_result['violence']
                    },
                    {
                        "tag_name": "sign",
                        "probability": check_result['sign']
                    },
                    {
                        "tag_name": "photo",
                        "probability": check_result['photo']
                    },
                    {
                        "tag_name": "fire",
                        "probability": check_result['fire']
                    },
                    {
                        "tag_name": "police",
                        "probability": check_result['police']
                    },
                    {
                        "tag_name": "children",
                        "probability": check_result['children']
                    },
                    {
                        "tag_name": "group_20",
                        "probability": check_result['group_20']
                    },
                    {
                        "tag_name": "group_100",
                        "probability": check_result['group_100']
                    },
                    {
                        "tag_name": "flag",
                        "probability": check_result['flag']
                    },
                    {
                        "tag_name": "night",
                        "probability": check_result['night']
                    },
                    {
                        "tag_name": "shouting",
                        "probability": check_result['shouting']
                    }]
            }
        }
        serializer.save(result=result)

        return Response(status=status.HTTP_201_CREATED)


class WordRecognitionViewSet(viewsets.ModelViewSet):

    queryset = WordRecognition.objects.all()
    serializer_class = WordRecognitionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()

        text = iserializer.text
        sensitive_list = sensitiveClass().check_sensitiveWords(text)

        if sensitive_list.get('sensitive_hit_flag') == 0:
            ret = 1
            msg = "无匹配记录"
        else:
            ret = 0
            msg = "匹配到记录"

        data = sensitive_list
        serializer.save(ret=ret, msg=msg, data=data, text=iserializer.text)

        return Response(status=status.HTTP_201_CREATED)


class WordRecognitionInspectionViewSet(viewsets.ModelViewSet):

    queryset = WordRecognitionInspection.objects.all()
    serializer_class = WordRecognitionInspectionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()

        # 增加网络URL文件上传
        if iserializer.text_url and not iserializer.text:
            txt_temp = NamedTemporaryFile(delete=True)
            txt_temp.write(urlopen(iserializer.text_url).read())
            txt_temp.flush()
            iserializer.text.save(os.path.basename(
                iserializer.text_url), File(txt_temp))

        # word格式文件读取
        filetype = FileType().filescanner(iserializer.text.path)
        if filetype is None:
            print('Cannot guess file type!')
            return Response(status=status.HTTP_400_BAD_REQUEST)

        text_content = ""
        sensitive_map = {}
        if filetype == 'zip':
            doc = docx.Document(iserializer.text.path)
            docText = '\n'.join(
                [paragraph.text for paragraph in doc.paragraphs])
            msg = "匹配记录"
            ret = 0
            text_content = docText
        elif filetype == 'wps':
            # 仅支持windows
            if(platform.system() == "Windows"):
                pythoncom.CoInitialize()
                word = wc.Dispatch("Word.Application")
                doc = word.Documents.Open(iserializer.text.path)
                docx_path = iserializer.text.path.split(".doc")[0] + '.docx'
                doc.SaveAs(docx_path, 12)
                doc.Close
                word.Quit
                doc = docx.Document(docx_path)
                docText = '\n'.join(
                    [paragraph.text for paragraph in doc.paragraphs])
                msg = "匹配记录"
                ret = 0
                text_content = docText
            else:
                # 仅支持ubuntu
                cmd = 'antiword -m UTF-8 ' + iserializer.text.path
                docText, errText = RunShellWithReturnCode(cmd)
                msg = "匹配记录"
                ret = 0
                if len(errText) > 0:
                    text_content = 'doc文档内容过短，请重新上传'
                else:
                    text_content = docText
        elif filetype == 'pdf':
            pdfText = PdfReader().parse(iserializer.text.path)
            msg = "匹配记录"
            ret = 0
            text_content = pdfText
        else:
            txtfile = iserializer.text.path
            # print(txtfile)
            # 增加gbk编码格式转换
            f_test = open(txtfile, 'rb')
            file_type = chardet.detect(f_test.read(100))

            if (file_type['encoding'] == 'GB2312'):
                f = codecs.open(txtfile, 'r', encoding='gbk', errors='ignore')
            elif (file_type['encoding'] == 'UTF-8-SIG'):
                f = codecs.open(txtfile, 'r', encoding='utf-8',
                                errors='ignore')
            elif (file_type['encoding'] == 'ascii'):
                f = codecs.open(txtfile, 'r', encoding='gbk', errors='ignore')
            else:
                f = codecs.open(txtfile, 'r', errors='ignore')

            try:
                for line in f:
                    text_content += line
            except Exception as e:
                print("The content get some error: " + line)
                print(e)
                msg = "内容获取异常"
                ret = 1
            else:
                print("Read content successfully!")
                msg = "匹配记录"
                ret = 0

        result = sensitiveClass().check_sensitiveWords(text_content)
        sensitive_map["text_content"] = text_content
        sensitive_map["sensitive_info"] = result
        data = sensitive_map
        serializer.save(ret=ret, msg=msg, data=result, text=iserializer.text)

        return Response(status=status.HTTP_201_CREATED)


class OcrGeneralViewSet(viewsets.ModelViewSet):

    queryset = OcrGeneral.objects.all()
    serializer_class = OcrGeneralSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "通用OCR"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataArr = []
        for each in arr:
            dataArr.append(each["text"])
        # result = check_result
        serializer.save(data=dataArr, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrIDCardViewSet(viewsets.ModelViewSet):

    queryset = OcrIDCard.objects.all()
    serializer_class = OcrIDCardSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "身份证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '姓名'):
                name = "name"
                count = count + 1
            if(each['name'] == '性别'):
                name = "sex"
                count = count + 1
            if(each['name'] == '民族'):
                name = "nation"
                count = count + 1
            if(each['name'] == '出生年月'):
                name = "birth"
                count = count + 1
            if(each['name'] == '身份证号码'):
                name = "id"
                count = count + 1
            if(each['name'] == '身份证地址'):
                name = "address"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 3):
            ret = 1
            msg = "请上传身份证图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class FileImageTerrorismUploadViewSet(viewsets.ModelViewSet):
    queryset = FileImageTerrorismUpload.objects.all()
    serializer_class = FileImageTerrorismUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        check_result = settings.VIOLENCE.check_violence(file_path)
        violence = check_result['violence']
        resultMap = {}
        resultMap['violence'] = get_two_float(float(violence) * 100, 2)
        serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)

        # 更新历史记录
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            resultMap, 'violence', violence, None)

        return Response(status=status.HTTP_201_CREATED)


class FileVisionPornUploadViewSet(viewsets.ModelViewSet):
    queryset = FileVisionPornUpload.objects.all()
    serializer_class = FileVisionPornUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # check_result = vision_porn(file_path)
        scores = settings.NSFW.caffe_preprocess_and_compute_api(file_path)
        resultMap = {}
        resultMap['normal_hot_porn'] = get_two_float(float(scores[1]) * 100, 2)
        # print (check_result)
        serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)

        # 更新历史记录
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            resultMap, 'porn', None, scores[1])

        return Response(status=status.HTTP_201_CREATED)


class VideoFileUploadViewSet(viewsets.ModelViewSet):
    queryset = VideoFileUpload.objects.all()
    serializer_class = VideoFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()

        # 增加网络URL文件上传
        if iserializer.video_url and not iserializer.video:
            video_temp = NamedTemporaryFile(delete=True)
            video_temp.write(urlopen(iserializer.video_url).read())
            video_temp.flush()
            iserializer.video.save(os.path.basename(
                iserializer.video_url), File(video_temp))

        file_path = iserializer.video.path
        resultMap = check_video(file_path)
        ret = 0
        msg = "成功"
        serializer.save(data=resultMap, ret=ret,
                        msg=msg, video=iserializer.video)
        return Response(status=status.HTTP_201_CREATED)


class AudioFileUploadViewSet(viewsets.ModelViewSet):
    queryset = AudioFileUpload.objects.all()
    serializer_class = AudioFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.speech_url and not iserializer.speech:
            speech_temp = NamedTemporaryFile(delete=True)
            speech_temp.write(urlopen(iserializer.speech_url).read())
            speech_temp.flush()
            iserializer.speech.save(os.path.basename(
                iserializer.speech_url), File(speech_temp))

        file_path = iserializer.speech.path
        size = os.path.getsize(file_path)
        if size <= 44:
            check_result = '录音时间太短，请重新录音！'
        else:
            check_result = audio().getOneAudioContent(file_path)
        # print (check_result)
        resultMap = {}
        resultMap['text'] = check_result
        serializer.save(data=resultMap, ret=ret, msg=msg,
                        speech=iserializer.speech)
        return Response(status=status.HTTP_201_CREATED)


class AudioFileInspectionViewSet(viewsets.ModelViewSet):
    queryset = AudioFileInspection.objects.all()
    serializer_class = AudioFileInspectionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.speech_url and not iserializer.speech:
            speech_temp = NamedTemporaryFile(delete=True)
            speech_temp.write(urlopen(iserializer.speech_url).read())
            speech_temp.flush()
            iserializer.speech.save(os.path.basename(
                iserializer.speech_url), File(speech_temp))

        # 转换mp3-wav
        kind = filetype.guess(iserializer.speech.path)
        if kind is None:
            print('Cannot guess file type!')
            return Response(status=status.HTTP_400_BAD_REQUEST)

        print(kind.extension)
        file_path = ''
        if kind.extension == 'mp3':
            sound = AudioSegment.from_mp3(iserializer.speech.path)
            destin_path = iserializer.speech.path.split(".mp3")[0] + '.wav'
            sound.export(destin_path, format='wav')
            file_path = destin_path
        elif kind.extension == 'ogg':
            sound = AudioSegment.from_ogg(iserializer.speech.path)
            destin_path = iserializer.speech.path.split(".ogg")[0] + '.wav'
            sound.export(destin_path, format='wav')
            file_path = destin_path
        else:
            file_path = iserializer.speech.path

        if len(file_path) > 0:
            duration = 0
            with contextlib.closing(wave.open(file_path, 'r')) as f:
                frames = f.getnframes()
                rate = f.getframerate()
                duration = frames / float(rate)
            audio_content = audio().getOneAudioContent(file_path)
            check_result = sensitiveClass().check_sensitiveWords(audio_content)
            resultMap = {}
            resultMap["speech_time"] = duration
            resultMap["speech_contents"] = check_result
            serializer.save(data=resultMap, ret=ret, msg=msg,
                            speech=iserializer.speech)
            return Response(status=status.HTTP_201_CREATED)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


class ImageFileUploadViewSet(viewsets.ModelViewSet):
    queryset = ImageFileUpload.objects.all()
    serializer_class = ImageFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        scores = settings.NSFW.caffe_preprocess_and_compute_api(file_path)
        resultMap = {}
        porn_sensitivity_level = "0"
        if (float(scores[1]) < 0.5):
            porn_sensitivity_level = "0"
        if (float(scores[1]) >= 0.5 and float(scores[1]) <= 0.9):
            porn_sensitivity_level = "1"
        if (float(scores[1]) > 0.9):
            porn_sensitivity_level = "2"
        resultMap['porn_sensitivity_level'] = porn_sensitivity_level
        resultMap['porn_percent'] = get_two_float(float(scores[1]) * 100, 2)

        check_result = settings.VIOLENCE.check_violence(file_path)
        violence = check_result['violence']
        violence_sensitivity_level = "0"
        if (float(violence) < 0.5):
            violence_sensitivity_level = "0"
        if (float(violence) >= 0.5 and float(violence) <= 0.9):
            violence_sensitivity_level = "1"
        if (float(violence) > 0.9):
            violence_sensitivity_level = "2"
        resultMap['violence_sensitivity_level'] = violence_sensitivity_level
        resultMap['violence_percent'] = get_two_float(float(violence) * 100, 2)

        resultMap['politics_sensitivity_level'] = ""
        resultMap['politics_percent'] = ""
        resultMap['public_character_level'] = ""
        resultMap['public_percent'] = ""

        serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)
        return Response(status=status.HTTP_201_CREATED)


class OcrDrivinglicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrDrivinglicense.objects.all()
    serializer_class = OcrDrivinglicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "驾驶证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '类型'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '证号'):
                name = "card_id"
                count = count + 1
            if(each['name'] == '姓名'):
                name = "driver"
                count = count + 1
            if(each['name'] == '性别'):
                name = "sex"
                count = count + 1
            if(each['name'] == '国籍'):
                name = "nationnality"
                count = count + 1
            if(each['name'] == '住址'):
                name = "address"
                count = count + 1
            if(each['name'] == '出生日期'):
                name = "birthday"
                count = count + 1
            if(each['name'] == '初次领证日期'):
                name = "issue_date"
                count = count + 1
            if(each['name'] == '准驾车型'):
                name = "class"
                count = count + 1
            if(each['name'] == '有效起始日期'):
                name = "valid_start"
                count = count + 1
            if(each['name'] == '有效截止日期'):
                name = "valid_end"
                count = count + 1

            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        # if (len(arr) == 0 or count < 1):
        if(dataMap["license_type"] != "中华人民共和国机动车驾驶证"):
            ret = 1
            msg = "请上传驾驶证图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrVehiclelicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrVehiclelicense.objects.all()
    serializer_class = OcrVehiclelicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "行驶证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '类型'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '证号'):
                name = "license_no"
                count = count + 1
            if(each['name'] == '号牌号码'):
                name = "plate_no"
                count = count + 1
            if(each['name'] == '车辆类型'):
                name = "vehicle_type"
                count = count + 1
            if(each['name'] == '所有人'):
                name = "owner"
                count = count + 1
            if(each['name'] == '住址'):
                name = "address"
                count = count + 1
            if(each['name'] == '使用性质'):
                name = "use_character"
                count = count + 1
            if(each['name'] == '品牌型号'):
                name = "model"
                count = count + 1
            if(each['name'] == '车辆识别代号'):
                name = "vin"
                count = count + 1
            if(each['name'] == '发动机号码'):
                name = "engine_no"
                count = count + 1
            if(each['name'] == '注册日期'):
                name = "register_date"
                count = count + 1
            if(each['name'] == '发证日期'):
                name = "issue_date"
                count = count + 1
            dataMap[name] = each['text']
            #dataMap[each['name']] = each['text']
        #result = check_result
        # if (len(arr) == 0 or count < 1):
        if(dataMap["license_type"] != "中华人民共和国机动车行驶证"):
            ret = 1
            msg = "请上传行驶证图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrBusinesslicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrBusinesslicense.objects.all()
    serializer_class = OcrBusinesslicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "营业执照"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '营业执照'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '统一社会信用代码'):
                name = "business_id"
                count = count + 1
            if(each['name'] == '名称'):
                name = "business_name"
                count = count + 1
            if(each['name'] == '类型'):
                name = "business_type"
                count = count + 1
            if(each['name'] == '住所'):
                name = "address"
                count = count + 1
            if(each['name'] == '法定代表人'):
                name = "operator"
                count = count + 1
            if(each['name'] == '注册资本'):
                name = "registered_capital"
                count = count + 1
            if(each['name'] == '成立日期'):
                name = "register_date"
                count = count + 1
            if(each['name'] == '营业期限'):
                name = "business_term"
                count = count + 1
            if(each['name'] == '经营范围'):
                name = "scope"
            dataMap[name] = each['text']
        if(dataMap["license_type"] != "营业执照"):
            ret = 1
            msg = "请上传营业执照图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrBankcardViewSet(viewsets.ModelViewSet):

    queryset = OcrBankcard.objects.all()
    serializer_class = OcrBankcardSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "银行卡"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '银行名称'):
                name = "bank_name"
                count = count + 1
            if(each['name'] == '卡号'):
                name = "bank_cardno"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 1):
            ret = 1
            msg = "请上传银行卡图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrHandWrittenViewSet(viewsets.ModelViewSet):

    queryset = OcrHandWritten.objects.all()
    serializer_class = OcrHandWrittenSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "手写体"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataArr = []
        for each in arr:
            dataArr.append(each["text"])
        # result = check_result
        serializer.save(data=dataArr, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class OcrVehicleplateViewSet(viewsets.ModelViewSet):

    queryset = OcrVehicleplate.objects.all()
    serializer_class = OcrVehicleplateSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "车牌"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        dataMap = {}
        count = 0
        for each in arr:
            name = ""
            if(each['name'] == '车牌'):
                name = "plate_no"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 1):
            ret = 1
            msg = "请上传车牌图片"
        serializer.save(data=dataMap, ret=ret, msg=msg,
                        image=iserializer.image)

        return Response(status=status.HTTP_201_CREATED)


class HistoryRecordViewSet(viewsets.ModelViewSet):

    queryset = HistoryRecord.objects.all()
    serializer_class = HistoryRecordSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def retrieve(self, request, pk=None):
        # 获取实例
        historyRecord = self.get_object()
        # 序列化
        serializer = self.get_serializer(historyRecord)
        dataMap = {}
        dataMap['ret'] = 0
        dataMap['msg'] = "成功"
        dataMap['data'] = serializer.data
        return Response(data=dataMap, status=status.HTTP_200_OK)

    def list(self, request):
        # 获取参数
        requestData = request.query_params
        system_id = requestData.get('system_id')
        channel_id = requestData.get('channel_id')
        user_id = requestData.get('user_id')
        begin_time = requestData.get('begin_time')
        end_time = requestData.get('end_time')
        file_name = requestData.get('file_name')
        file_type = requestData.get('file_type')
        is_group = requestData.get('is_group')
        query_date = requestData.get('query_date')

        # 根据条件过滤
        conditions = {}
        if system_id is not None:
            conditions['system_id'] = system_id

        if channel_id is not None:
            conditions['channel_id'] = channel_id

        if user_id is not None:
            conditions['user_id'] = user_id

        if file_name is not None:
            conditions['file_name'] = file_name

        if file_type is not None:
            conditions['file_type'] = file_type

        if begin_time is not None:
            begin_time_date = datetime.datetime.strptime(
                begin_time, "%Y-%m-%d %H:%M:%S")
            conditions['upload_time__gte'] = begin_time_date

        if end_time is not None:
            end_time_date = datetime.datetime.strptime(
                end_time, "%Y-%m-%d %H:%M:%S")
            conditions['upload_time__lte'] = end_time_date

        if is_group is not None and is_group == 'true':
            historygroups = HistoryRecord.objects.filter().extra(
                {'day': "strftime('%Y-%m-%d',upload_time)"}).values_list('day').annotate(Count('id')).order_by('-day')
            print(historygroups)
            result = {}
            results = {}
            line = []
            if len(historygroups) > 0:
                if query_date is not None and int(query_date) > 0:
                    group_index = 0
                    for historygroup in historygroups:
                        if group_index < int(query_date):
                            historydate = historygroup[0]
                            print(historydate)
                            conditions["upload_time__year"] = historydate.split(
                                '-')[0]
                            conditions["upload_time__month"] = historydate.split(
                                '-')[1]
                            conditions["upload_time__day"] = historydate.split(
                                '-')[2]
                            historylist = HistoryRecord.objects.filter(
                                **conditions)
                            serializer_group = self.get_serializer(
                                historylist, many=True)
                            result[historydate] = serializer_group.data
                            group_index += 1
                        else:
                            results['results'] = result
                            return Response(results)
                else:
                    for historygroup in historygroups:
                        historydate = historygroup[0]
                        print(historydate)
                        conditions["upload_time__year"] = historydate.split(
                            '-')[0]
                        conditions["upload_time__month"] = historydate.split(
                            '-')[1]
                        conditions["upload_time__day"] = historydate.split(
                            '-')[2]
                        historylist = HistoryRecord.objects.filter(
                            **conditions)
                        serializer_group = self.get_serializer(
                            historylist, many=True)
                        result[historydate] = serializer_group.data
                    results['results'] = result
                    return Response(results)
            else:
                return Response([])
        else:
            queryset = HistoryRecord.objects.filter(**conditions)

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
